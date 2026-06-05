import os
import re
import io
import sys
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def add_page_number(run):
    """Inserts a PAGE field into the run for dynamic word page numbers."""
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    r = run._r
    r.append(fldChar1)
    r.append(instrText)
    r.append(fldChar2)
    r.append(fldChar3)

def tokenize_python_code(code_text):
    """
    Splits Python code into tokens of (type, text) for custom syntax highlighting.
    Robustly handles keywords, built-ins, comments, and strings.
    """
    keywords = {
        'def', 'class', 'return', 'if', 'else', 'elif', 'for', 'while', 
        'try', 'except', 'finally', 'import', 'from', 'as', 'with', 'in', 
        'is', 'not', 'and', 'or', 'pass', 'break', 'continue', 'lambda', 
        'global', 'nonlocal', 'assert', 'yield', 'del', 'None', 'True', 'False'
    }
    builtins = {
        'print', 'len', 'sum', 'min', 'max', 'abs', 'round', 'int', 'float', 
        'str', 'list', 'dict', 'set', 'tuple', 'zip', 'enumerate', 'range', 
        'any', 'all', 'isinstance', 'type', 'open', 'super', 'classmethod', 'object', 'setattr', 'getattr', 'hasattr'
    }
    
    # Combined regex pattern
    pattern = re.compile(
        r'(?P<COMMENT>#.*)'
        r'|(?P<STRING_TRIPLE_DQ>"""[\s\S]*?""")'
        r'|(?P<STRING_TRIPLE_SQ>\'\'\'[\s\S]*?\'\'\')'
        r'|(?P<STRING_DQ>"[^"\\]*(?:\\.[^"\\]*)*")'
        r'|(?P<STRING_SQ>\'[^\'\\]*(?:\\.[^\'\\]*)*\')'
        r'|(?P<NUMBER>\b\d+(?:\.\d*)?\b)'
        r'|(?P<NAME>[a-zA-Z_][a-zA-Z0-9_]*)'
        r'|(?P<SPACE>[ \t]+)'
        r'|(?P<NEWLINE>\n)'
        r'|(?P<ANY>.)'
    )
    
    tokens = []
    for m in pattern.finditer(code_text):
        token_type = m.lastgroup
        token_value = m.group(token_type)
        
        if token_type == 'NAME':
            if token_value in keywords:
                type_ = 'KEYWORD'
            elif token_value in builtins:
                type_ = 'BUILTIN'
            elif token_value in ('self', 'cls'):
                type_ = 'SELF'
            else:
                type_ = 'NAME'
        elif token_type in ('STRING_TRIPLE_DQ', 'STRING_TRIPLE_SQ', 'STRING_DQ', 'STRING_SQ'):
            type_ = 'STRING'
        elif token_type in ('COMMENT', 'NUMBER', 'SPACE', 'NEWLINE'):
            type_ = token_type
        else:
            type_ = 'PUNCTUATION'
            
        tokens.append((type_, token_value))
    return tokens

def split_multiline_tokens(tokens):
    """Splits tokens containing newlines into distinct segments per line."""
    flat_tokens = []
    for type_, val in tokens:
        if '\n' in val and type_ != 'NEWLINE':
            lines = val.split('\n')
            for i, line_val in enumerate(lines):
                if i > 0:
                    flat_tokens.append(('NEWLINE', '\n'))
                if line_val:
                    flat_tokens.append((type_, line_val))
        else:
            flat_tokens.append((type_, val))
    return flat_tokens

def group_tokens_by_line(tokens):
    """Groups flat tokens list into a list of lines, each line a list of tokens."""
    lines = []
    current_line = []
    for type_, val in tokens:
        if type_ == 'NEWLINE':
            lines.append(current_line)
            current_line = []
        else:
            current_line.append((type_, val))
    if current_line or not lines:
        lines.append(current_line)
    return lines

def format_code_block(doc, title, filepath, code_text, description):
    """Creates a beautifully styled, syntax-highlighted code section inside a table block."""
    # Heading 2
    h2 = doc.add_heading(title, level=2)
    h2.paragraph_format.space_before = Pt(20)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.keep_with_next = True
    
    # Description
    p_desc = doc.add_paragraph()
    p_desc.paragraph_format.space_after = Pt(12)
    p_desc.paragraph_format.line_spacing = 1.15
    run_desc = p_desc.add_run(description)
    run_desc.font.name = 'Times New Roman'
    run_desc.font.size = Pt(11)
    
    # Read and parse code
    tokens = tokenize_python_code(code_text)
    flat_tokens = split_multiline_tokens(tokens)
    lines_tokens = group_tokens_by_line(flat_tokens)
    
    # Width and page alignment setup for table
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.autofit = False
    
    # Cell styling
    cell = table.cell(0, 0)
    # 6.5 inches printable width (8.5 - 2 inches margins)
    cell.width = Inches(6.5)
    
    # Cell background shading (Light gray #F8F9FA)
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="F8F9FA"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))
    
    # Cell thin borders (Muted gray #D3D3D3)
    borders_xml = f'''
    <w:tcBorders {nsdecls("w")}>
        <w:top w:val="single" w:sz="4" w:space="0" w:color="D3D3D3"/>
        <w:left w:val="single" w:sz="4" w:space="0" w:color="D3D3D3"/>
        <w:bottom w:val="single" w:sz="4" w:space="0" w:color="D3D3D3"/>
        <w:right w:val="single" w:sz="4" w:space="0" w:color="D3D3D3"/>
    </w:tcBorders>
    '''
    cell._tc.get_or_add_tcPr().append(parse_xml(borders_xml))
    
    # Line number width formatting
    total_lines = len(lines_tokens)
    ln_width = len(str(total_lines))
    
    for idx, line_toks in enumerate(lines_tokens, 1):
        p = cell.paragraphs[0] if idx == 1 else cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.05
        
        # 1. Add line number (gray, monospaced)
        ln_run = p.add_run(f"{idx:>{ln_width}}  ")
        ln_run.font.name = 'Consolas'
        ln_run.font.size = Pt(8.5)
        ln_run.font.color.rgb = RGBColor(140, 140, 140)
        
        # 2. Add tokens with custom syntax highlighting
        for type_, val in line_toks:
            r = p.add_run(val)
            r.font.name = 'Consolas'
            r.font.size = Pt(9.0)
            
            if type_ == 'KEYWORD':
                r.font.color.rgb = RGBColor(0, 51, 204)  # Deep Blue
                r.font.bold = True
            elif type_ == 'BUILTIN':
                r.font.color.rgb = RGBColor(0, 128, 128)  # Teal
            elif type_ == 'SELF':
                r.font.color.rgb = RGBColor(160, 80, 0)   # Brown
                r.font.italic = True
            elif type_ == 'COMMENT':
                r.font.color.rgb = RGBColor(46, 125, 50)  # Forest Green
                r.font.italic = True
            elif type_ == 'STRING':
                r.font.color.rgb = RGBColor(163, 21, 21)  # Dark Red/Rust
            elif type_ == 'NUMBER':
                r.font.color.rgb = RGBColor(106, 27, 154) # Purple
            else:
                r.font.color.rgb = RGBColor(0, 0, 0)      # Black (Default)

def build_word_document():
    doc = Document()
    
    # 1. Page Setup (Letter size, 1.0 inch margins)
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11.0)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
        # Enable different first page if needed (not active here, keeping simple)
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False
        
        # Header setup
        header_p = section.header.paragraphs[0]
        header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        header_run = header_p.add_run("Appendix B: Recommendation Engine Source Code")
        header_run.font.name = 'Times New Roman'
        header_run.font.size = Pt(8.5)
        header_run.font.italic = True
        header_run.font.color.rgb = RGBColor(128, 128, 128)
        
        # Footer setup
        footer_p = section.footer.paragraphs[0]
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer_p.add_run("Page ")
        footer_run.font.name = 'Times New Roman'
        footer_run.font.size = Pt(10)
        add_page_number(footer_run)
        
    # Configure document base style
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Times New Roman'
    style_normal.font.size = Pt(11)
    
    # Configure Headings
    style_h1 = doc.styles['Heading 1']
    style_h1.font.name = 'Times New Roman'
    style_h1.font.size = Pt(18)
    style_h1.font.bold = True
    style_h1.font.color.rgb = RGBColor(0, 0, 0)
    
    style_h2 = doc.styles['Heading 2']
    style_h2.font.name = 'Times New Roman'
    style_h2.font.size = Pt(13)
    style_h2.font.bold = True
    style_h2.font.color.rgb = RGBColor(0, 0, 0)
    
    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(24)
    title_p.paragraph_format.space_after = Pt(12)
    title_run = title_p.add_run("APPENDIX B: SOURCE CODE")
    title_run.font.name = 'Times New Roman'
    title_run.font.size = Pt(20)
    title_run.font.bold = True
    
    # Subtitle or explanation
    intro_p = doc.add_paragraph()
    intro_p.paragraph_format.space_after = Pt(24)
    intro_p.paragraph_format.line_spacing = 1.15
    intro_run = intro_p.add_run(
        "This appendix provides the complete, production-ready source code for the core machine learning "
        "and recommendation engine modules of the application. To maintain document brevity and focus on "
        "algorithmic substance, only the most critical algorithmic scripts (Collaborative Filtering, "
        "Content-Based Filtering, Hybrid Engine, and Serving/Blending Layer) are included. "
        "These modules together drive the personalized recommendation features, session-based recommendation for anonymous visitors, "
        "real-time caching, and product category diversity enforcement."
    )
    intro_run.font.name = 'Times New Roman'
    intro_run.font.size = Pt(11)
    
    # Read files and compile them
    engine_dir = os.path.join("recommendation", "engine")
    
    # 1. Collaborative Filtering
    collab_path = os.path.join(engine_dir, "collaborative.py")
    with open(collab_path, "r", encoding="utf-8") as f:
        collab_code = f.read()
    collab_desc = (
        "This module implements the Collaborative Filtering recommendation engine using a Latent Factor Model based on "
        "Truncated Singular Value Decomposition (SVD). SVD decomposes the sparse user-item interaction matrix into two lower-rank "
        "matrices representing user and item latent factors. Ratings are normalized by subtracting the average rating of each "
        "user to account for rating bias, and the Truncated SVD model fits the centered matrix to predict ratings via dot product "
        "of latent vectors, which are then clipped to a valid 1.0-5.0 range."
    )
    format_code_block(
        doc=doc,
        title="B.1 Collaborative Filtering Algorithm (collaborative.py)",
        filepath=collab_path,
        code_text=collab_code,
        description=collab_desc
    )
    
    # 2. Content Based Filtering
    content_path = os.path.join(engine_dir, "content_based.py")
    with open(content_path, "r", encoding="utf-8") as f:
        content_code = f.read()
    content_desc = (
        "This module implements the Content-Based Filtering engine using TF-IDF (Term Frequency-Inverse Document Frequency) text "
        "vectorization and Cosine Similarity. Product metadata is parsed to build high-dimensional TF-IDF vectors, and user profile "
        "vectors are created as weighted averages of their interactions (views get weight 1.0, purchases/high ratings get 3.0, "
        "and low ratings get negative weights). These profiles are then matched with all catalog items using cosine similarity "
        "to yield semantically relevant recommendations."
    )
    format_code_block(
        doc=doc,
        title="B.2 Content-Based Filtering Algorithm (content_based.py)",
        filepath=content_path,
        code_text=content_code,
        description=content_desc
    )
    
    # 3. Hybrid Engine
    hybrid_path = os.path.join(engine_dir, "hybrid.py")
    with open(hybrid_path, "r", encoding="utf-8") as f:
        hybrid_code = f.read()
    hybrid_desc = (
        "This module implements the Hybrid Recommendation Engine, which dynamically blends Collaborative Filtering (CF) and "
        "Content-Based Filtering (CBF) predictions based on each user's history density. The mixing coefficient (alpha) increases "
        "proportionally with the user's total ratings (capped at 20 ratings for alpha=1.0). High-density users leverage community "
        "signals (CF), while cold-start or low-activity users leverage item attributes (CBF). Blending scores are normalized using "
        "Min-Max scaling and boosted with real-time search keywords."
    )
    format_code_block(
        doc=doc,
        title="B.3 Hybrid Recommendation Engine (hybrid.py)",
        filepath=hybrid_path,
        code_text=hybrid_code,
        description=hybrid_desc
    )
    
    # 4. Serving
    serving_path = os.path.join(engine_dir, "serving.py")
    with open(serving_path, "r", encoding="utf-8") as f:
        serving_code = f.read()
    serving_desc = (
        "This module implements the serving and real-time blending layer of the recommendation system. It bridges the Django application "
        "views with the underlying machine learning models. The serving layer features an in-memory caching system to avoid expensive disk "
        "read operations, interleaves pre-computed recommendations with real-time user-interest signals to maintain freshness, enforces "
        "a category diversity booster that caps items at two per category, and provides anonymous guest tracking via session keys."
    )
    format_code_block(
        doc=doc,
        title="B.4 Recommendation Serving Layer (serving.py)",
        filepath=serving_path,
        code_text=serving_code,
        description=serving_desc
    )
    
    # Save the output documents
    # 1. Inside ecommerce_recs
    out_path_1 = "Appendix_B_Source_Code.docx"
    doc.save(out_path_1)
    print(f"Successfully saved docx to {out_path_1}")
    
    # 2. Inside documentation (main documentation folder)
    out_path_2 = os.path.join("..", "documentation", "Appendix_B_Source_Code.docx")
    try:
        doc.save(out_path_2)
        print(f"Successfully saved docx to {out_path_2}")
    except Exception as e:
        print(f"Could not save to {out_path_2} (is it open in Word?): {e}")

if __name__ == "__main__":
    build_word_document()
